"""D8: the docx, with its provenance carried inside the document.

The provenance map is not an API nicety bolted on beside the file — it ships in the file.
Every bullet in the Experience section carries a bracketed reference, and the final page
resolves each reference to the exact fact statements that license the sentence, plus every
omission with its typed reason. A printed copy still traces every sentence to a fact.
"""
from __future__ import annotations

import io

from docx import Document
from docx.shared import Pt

from .facts import AtomicClaim
from .selector import Omission


def render_docx(
    candidate_name: str,
    job_title: str,
    bullets: list[dict],
    claims_by_id: dict[str, AtomicClaim],
    omissions: list[Omission],
    entailment_note: str,
) -> bytes:
    """bullets: [{position, text, cites, entailment}] — already gated, never raw drafts."""
    doc = Document()

    doc.add_heading(candidate_name, level=0)
    sub = doc.add_paragraph()
    run = sub.add_run(f"Compiled for: {job_title}")
    run.italic = True

    doc.add_heading("Experience", level=1)
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(b["text"])
        ref = p.add_run(f"  [{b['position'] + 1}]")
        ref.font.size = Pt(8)

    doc.add_page_break()
    doc.add_heading("Provenance", level=1)
    doc.add_paragraph(
        "Every sentence above cites the career facts that license it. The compiler rejects "
        "any sentence that cites nothing, cites a fact that does not exist, carries a number "
        "appearing in no cited fact, or claims more than its cited evidence supports "
        f"({entailment_note}).")

    for b in bullets:
        head = doc.add_paragraph()
        head.add_run(f"[{b['position'] + 1}] ").bold = True
        head.add_run(b["text"])
        for cid in b["cites"]:
            claim = claims_by_id.get(cid)
            line = doc.add_paragraph(style="List Bullet 2")
            if claim is None:
                line.add_run(f"{cid} (not resolved)")
                continue
            line.add_run(f"{claim.claim_key}: ").bold = True
            line.add_run(claim.core.statement)
            meta = line.add_run(f"  — {claim.provenance.value}, fact {cid}")
            meta.font.size = Pt(8)
        score = b.get("entailment")
        if score is not None:
            s = doc.add_paragraph(style="List Bullet 2")
            s.add_run(f"entailment {score:.2f}").font.size = Pt(8)

    if omissions:
        doc.add_heading("What was left out, and why", level=1)
        doc.add_paragraph(
            "A selector that drops evidence silently is not explainable. Every eligible fact "
            "that did not reach the page:")
        for o in omissions:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{o.claim_key}: ").bold = True
            p.add_run(f"{o.reason.value} — {o.detail}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def letter_scaffold(candidate_name: str, job_title: str) -> dict:
    """The deterministic letter frame. Template text owned by code, never the model:
    these are the only sentences allowed to reference the role directly, and they can
    fabricate nothing about the candidate because they say nothing about the candidate."""
    return {
        "greeting": "Dear Hiring Team,",
        "opening": f"I am writing to apply for the {job_title} role. "
                   "Every statement below cites a verified career fact; "
                   "the compiler rejected anything stronger than its evidence.",
        "closing": "I would welcome the chance to discuss any of these points, "
                   "each of which traces to its source.",
        "signoff": f"Sincerely,\n{candidate_name}",
    }


def render_letter_docx(
    candidate_name: str,
    job_title: str,
    sentences: list[dict],
    claims_by_id: dict[str, AtomicClaim],
    entailment_note: str,
) -> bytes:
    """sentences: [{position, text, cites, entailment}] — already gated, never raw drafts.
    The provenance map ships inside the document, same as the resume."""
    frame = letter_scaffold(candidate_name, job_title)
    doc = Document()

    doc.add_heading(candidate_name, level=0)
    sub = doc.add_paragraph()
    sub.add_run(f"Cover letter — {job_title}").italic = True

    doc.add_paragraph(frame["greeting"])
    doc.add_paragraph(frame["opening"])
    for s in sentences:
        p = doc.add_paragraph()
        p.add_run(s["text"])
        ref = p.add_run(f"  [{s['position'] + 1}]")
        ref.font.size = Pt(8)
    doc.add_paragraph(frame["closing"])
    doc.add_paragraph(frame["signoff"])

    doc.add_page_break()
    doc.add_heading("Provenance", level=1)
    doc.add_paragraph(
        "Every bracketed sentence above cites the career facts that license it. The "
        "compiler rejects any sentence that cites nothing, cites a fact that does not "
        "exist, carries a number appearing in no cited fact, or claims more than its "
        f"cited evidence supports ({entailment_note}). The greeting, the role line, and "
        "the closing are fixed template text: they reference the role, and say nothing "
        "about the candidate.")
    for s in sentences:
        head = doc.add_paragraph()
        head.add_run(f"[{s['position'] + 1}] ").bold = True
        head.add_run(s["text"])
        for cid in s["cites"]:
            claim = claims_by_id.get(cid)
            line = doc.add_paragraph(style="List Bullet 2")
            if claim is None:
                line.add_run(f"{cid} (not resolved)")
                continue
            line.add_run(f"{claim.claim_key}: ").bold = True
            line.add_run(claim.core.statement)
            meta = line.add_run(f"  — {claim.provenance.value}, fact {cid}")
            meta.font.size = Pt(8)
        score = s.get("entailment")
        if score is not None:
            sc = doc.add_paragraph(style="List Bullet 2")
            sc.add_run(f"entailment {score:.2f}").font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
